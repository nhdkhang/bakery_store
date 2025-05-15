from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify 
from flask_mysqldb import MySQL
from flask_mail import Mail, Message
from MySQLdb.cursors import DictCursor
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from functools import wraps
from datetime import datetime, timedelta, time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import time
import traceback
import json
import random
import requests
import uuid
import hmac
import hashlib
import random
import smtplib


app = Flask(__name__)
mysql = MySQL(app)

app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = '123456'
app.config['MYSQL_DB'] = 'HomemadeBakery'

app.config['SECRET_KEY'] = 'VN45dscdf32raafeqqq38ad_fqe1@#sdas'
app.config['SESSION_PERMANENT'] = True
app.config['SESSION_TYPE'] = 'filesystem'


app.config.update(
    MAIL_SERVER='smtp.gmail.com',
    MAIL_PORT=587,
    MAIL_USE_TLS=True,
    MAIL_USE_SSL=False,
    MAIL_USERNAME='homemade.bakery04301975@gmail.com',
    MAIL_PASSWORD='umfc cgiq nyfj raag',
    MAIL_DEFAULT_SENDER=('HomeMade Bakery', 'homemade.bakery04301975@gmail.com')
)

mail = Mail(app)

UPLOAD_FOLDER = 'static/uploads/payments'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

#------------------------------------------DECORATORS----------------------------------------------
def calculate_total_price(cart_items):
    total_price = 0
    for item in cart_items:
        total_price += item['price'] * item['quantity']
    return total_price

def generate_otp():
    otp = ''.join([str(random.randint(0, 9)) for _ in range(6)])
    expiration = datetime.now() + timedelta(minutes=5)
    return otp, expiration

def get_product_image(product_id):
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("SELECT image_url FROM Products WHERE product_id = %s", (product_id,))
        product = cursor.fetchone()
        return product['image_url'] if product else None
    except Exception as e:
        print(f"Error getting product image: {str(e)}")
        return None
    finally:
        cursor.close()

def get_active_promotions(customer_id=None):
    """Get active promotions that are valid for the current date and haven't been used by the customer"""
    cursor = mysql.connection.cursor(DictCursor)
    try:
        query = """
            SELECT
                p.promotion_id,
                p.promotion_code,
                p.promotion_name,
                p.discount_percentage,
                p.start_date,
                p.end_date,
                p.max_uses,
                p.current_uses
            FROM Promotions p
            WHERE CURDATE() BETWEEN p.start_date AND p.end_date
            AND (p.max_uses = 0 OR p.current_uses < p.max_uses)
        """
        params = []
        
        if customer_id:
            query += """
                AND NOT EXISTS (
                    SELECT 1 
                    FROM Orders o 
                    WHERE o.promotion_id = p.promotion_id
                    AND o.customer_id = %s
                    AND o.status != 'cancelled'
                )
            """
            params.append(customer_id)
            
        query += " ORDER BY p.discount_percentage DESC"
        
        cursor.execute(query, tuple(params))
        return cursor.fetchall()
    except Exception as e:
        print(f"Error getting active promotions: {str(e)}")
        return []
    finally:
        cursor.close()

def send_otp_email(email):
    try:
        otp, expiration = generate_otp()
        msg = Message(
            subject="HomeMade Bakery - Your OTP Code",
            sender=app.config['MAIL_DEFAULT_SENDER'],
            recipients=[email]
        )
        
        msg.body = f"""
        Your OTP code is: {otp}
        
        This code will expire in 5 minutes.
        If you didn't request this code, please ignore this email.
        
        Best regards,
        HomeMade Bakery Team
        """
        mail.send(msg)
        print(f"OTP email sent successfully to {email}")
        return otp, expiration
    except Exception as e:
        print(f"Error sending OTP email: {str(e)}")
        print(f"Error details: {traceback.format_exc()}")
        return None, None
    
@app.route('/test-mail')
def test_mail():
    try:
        msg = Message(
            subject='Test Email from HomeMade Bakery',
            recipients=['your-test-email@example.com'],
            body='This is a test email to verify SMTP configuration'
        )
        mail.send(msg)
        return 'Test email sent successfully!'
    except Exception as e:
        return f'Error sending email: {str(e)}\n{traceback.format_exc()}'
#------------------------------------------HOME--------------------------------------------------
@app.route('/')
def home():
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("""
            SELECT p.*, c.category_name 
            FROM Products p
            JOIN Categories c ON p.category_id = c.category_id
            WHERE p.product_id IN (
                SELECT MIN(product_id)
                FROM Products
                GROUP BY category_id
            )
        """)
        featured_products = cursor.fetchall()
        
        cursor.execute(""" 
            SELECT r.*, p.product_name, c.fullname as customer_name
            FROM Reviews r
            JOIN Products p ON r.product_id = p.product_id
            JOIN Customers c ON r.customer_id = c.customer_id
            ORDER BY r.created_at DESC
            LIMIT 3
        """)
        reviews = cursor.fetchall()

        for review in reviews:
            review['product_image_url'] = get_product_image(review['product_id'])
        
        return render_template('home.html', random_products=featured_products, reviews=reviews)
    except Exception as e:
        flash(f'Error loading products: {str(e)}', 'danger')
        return render_template('home.html', random_products=[])
    finally:
        cursor.close()

#------------------------------------------AUTH--------------------------------------------------
@app.route('/register', methods=['GET', 'POST']) 
def register():
    if request.method == 'POST':
        email = request.form['email']
        fullname = request.form['fullname']
        phone_number = request.form['phone_number']
        address = request.form['address']
        password = generate_password_hash(request.form['password'])

        if not email or not password or not fullname or not phone_number or not address:
            flash('Please fill in all fields.', 'danger')
            return render_template('register.html')

        if len(phone_number) > 10 or not phone_number.isdigit():
            flash('Phone number is not valid.', 'danger')
            return render_template('register.html')

        try:
            cursor = mysql.connection.cursor()
            cursor.execute("INSERT INTO Users (email, password, role) VALUES (%s, %s, 'customer')", (email, password))
            user_id = cursor.lastrowid
            cursor.execute("INSERT INTO Customers (fullname, phone_number, address, user_id) VALUES (%s, %s, %s, %s)",
                           (fullname, phone_number, address, user_id))
            mysql.connection.commit()
            flash('Registered successfully!', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            mysql.connection.rollback()
            flash(f'Error: {e}', 'danger')

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form.get('email')
        password = request.form.get('password')
        

        if not email or not password:
            flash('Please provide both email and password', 'danger')
            return render_template('login.html')

        try:
            cursor = mysql.connection.cursor(DictCursor)
            cursor.execute("SELECT * FROM Users WHERE email = %s", (email,))
            user = cursor.fetchone()

            if user and check_password_hash(user['password'], password):
                if user['role'] == 'customer':
                    otp, expiration = send_otp_email(email)
                    if otp and expiration:
                        session['temp_user_id'] = user['user_id']
                        session['otp'] = otp
                        session['otp_expiration'] = expiration.isoformat()
                        flash('OTP has been sent to your email', 'success')
                        return redirect(url_for('verify_otp'))
                    else:
                        flash('Failed to send OTP. Please try again.', 'danger')
                        return redirect(url_for('login'))
                else:
                    session['user_id'] = user['user_id']
                    session['role'] = user['role'] 
                    flash('Login successful as administrator!', 'success')
                return redirect(url_for('admin_dashboard'))
            else:
                flash('Invalid email or password', 'danger')

        except Exception as e:
            flash(f'Error: {str(e)}', 'danger')
        finally:
            if cursor:
                cursor.close()
    return render_template('login.html')

@app.route('/verify_otp', methods=['GET', 'POST'])
def verify_otp():
    if request.method == 'POST':
        user_otp = request.form.get('otp')
        stored_otp = session.get('otp')
        otp_expiration = session.get('otp_expiration')
        
        if not user_otp or not stored_otp or not otp_expiration:
            flash('Invalid OTP request. Please try again.', 'danger')
            return redirect(url_for('login'))
        
        if datetime.now() > datetime.fromisoformat(otp_expiration):
            session.pop('otp', None)
            session.pop('otp_expiration', None)
            flash('OTP has expired. Please request a new one.', 'danger')
            return redirect(url_for('login'))
        
        if user_otp == stored_otp:
            session.pop('otp', None)
            session.pop('otp_expiration', None)
            
            user_id = session.pop('temp_user_id', None)
            if user_id:
                session['user_id'] = user_id
                cursor = mysql.connection.cursor(DictCursor)
                cursor.execute("SELECT role FROM Users WHERE user_id = %s", (user_id,))
                user = cursor.fetchone()
                session['role'] = user['role']
                cursor.close()
                
                flash('Login successful!', 'success')
                if user['role'] == 'customer':
                    return redirect(url_for('customer_dashboard'))
                else:
                    return redirect(url_for('admin_dashboard'))
        else:
            flash('Invalid OTP. Please try again.', 'danger')
    
    return render_template('verify_otp.html')

@app.route('/send-otp')
def send_otp():
    try:
        msg = Message(
            subject='Mã OTP của bạn',
            sender='homemade.bakery04301975@gmail.com',
            recipients=['recipient@example.com'],
            body=f'Your One Time Password is: 123456'
        )
        mail.send(msg)
        return "OTP is sent successfully!"
    except Exception as e:
        return f"Error: {str(e)}"

@app.route('/logout')
def logout():
    session.clear()
    flash('Logout successful!', 'success')
    return redirect(url_for('home'))
#-----------------------------------------PRODUCTS------------------------------------------------
@app.route('/products')
def products():
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute(""" 
        SELECT p.*, c.category_name 
        FROM Products p
        JOIN Categories c ON p.category_id = c.category_id
        """)
        products = cursor.fetchall()

        cursor.execute("SELECT * FROM Categories")
        categories = cursor.fetchall()
        return render_template('products.html', products=products, categories=categories)
    except Exception as e:
        flash(f'Error loading products: {str(e)}', 'danger')
        return render_template('products.html', products=[], categories=[])
    finally:
        cursor.close()
#------------------------------------------DASHBOARDS----------------------------------------------
@app.route('/customer_dashboard', methods=['GET'])
def customer_dashboard():
    if 'user_id' in session and session.get('role') == 'customer':
        try:
            cursor = mysql.connection.cursor(DictCursor)
            cursor.execute("""
                SELECT c.customer_id, fullname 
                FROM Customers c
                WHERE user_id = %s
            """, (session['user_id'],))
            customer = cursor.fetchone()
            if not customer:
                flash('Customer information not found.', 'danger')
                return redirect(url_for('home'))
            else:
                session['customer_id'] = customer['customer_id']
                customer_fullname = customer['fullname']


            query = request.args.get('query', '')
            category_id = request.args.get('category_id', '')
            
            sql = """
                SELECT p.*, c.category_name 
                FROM Products p
                JOIN Categories c ON p.category_id = c.category_id
                WHERE 1=1
            """
            params = []

            if query:
                sql += " AND p.product_name LIKE %s"
                params.append(f'%{query}%')
            
            if category_id:
                sql += " AND p.category_id = %s"
                params.append(category_id)
            
            cursor.execute(sql, tuple(params))
            products = cursor.fetchall()
            
            cursor.execute("SELECT * FROM Categories")
            categories = cursor.fetchall()
            cursor.close()
            
            return render_template(
                'customer_dashboard.html',
                customer_fullname=customer_fullname,
                products=products,
                categories=categories,
                query=query,
                selected_category_id=category_id
            )
        
        except Exception as e:
            flash(f'Lỗi: {e}', 'danger')
            return redirect(url_for('home'))
    else:
        return redirect(url_for('login'))

@app.route('/admin_dashboard')
def admin_dashboard():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    return render_template('admin_dashboard.html')
#-----------------------------------------SEARCH-------------------------------------------------
@app.route('/search', methods=['GET'])
def search():
    query = request.args.get('query', '')
    category_id = request.args.get('category_id', '')
    min_price = request.args.get('min_price', '')
    max_price = request.args.get('max_price', '')
    sort_by = request.args.get('sort_by', '')
    referrer = request.referrer

    try:
        cursor = mysql.connection.cursor(DictCursor)

        sql = """
            SELECT p.*, c.category_name 
            FROM Products p
            JOIN Categories c ON p.category_id = c.category_id
            WHERE 1=1
        """
        params = []

        if query:
            sql += " AND p.product_name LIKE %s"
            params.append(f'%{query}%')
        
        if category_id:
            sql += " AND p.category_id = %s"
            params.append(category_id)
        
        if min_price:
            sql += " AND p.price >= %s"
            params.append(float(min_price))
        
        if max_price:
            sql += " AND p.price <= %s"
            params.append(float(max_price))

        sort_options = {
            'price_asc': " ORDER BY p.price ASC",
            'price_desc': " ORDER BY p.price DESC",
            'name_asc': " ORDER BY p.product_name ASC",
            'name_desc': " ORDER BY p.product_name DESC"
        }
        sql += sort_options.get(sort_by, '')

        cursor.execute(sql, tuple(params))
        products = cursor.fetchall()

        cursor.execute("SELECT * FROM Categories")
        categories = cursor.fetchall()

        if referrer and 'customer_dashboard' in referrer:
            template = 'customer_dashboard.html'
        else:
            template = 'products.html'
            
        return render_template(
            template,
            products=products,
            categories=categories,
            query=query,
            selected_category_id=category_id,
            min_price=min_price,
            max_price=max_price,
            sort_by=sort_by
        )
        
    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))
    
    finally:
        if cursor:
            cursor.close()

@app.route('/product_detail/<int:product_id>', methods=['GET', 'POST'])
def product_detail(product_id):
    if request.method == 'POST':
        if 'user_id' not in session or session.get('role') != 'customer':
            flash('You need to login to add items to your cart.', 'danger')
            return redirect(url_for('login'))

        try:
            cursor = mysql.connection.cursor(DictCursor)
            cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
            customer = cursor.fetchone()

            if not customer:
                flash('Customer information is not available.', 'danger')
                return redirect(url_for('customer_dashboard'))

            cursor.execute("SELECT cart_id FROM Cart WHERE customer_id = %s AND cart_status = 'pending'", (customer['customer_id'],))
            cart = cursor.fetchone()

            if not cart:
                cursor.execute("INSERT INTO Cart (customer_id, cart_status) VALUES (%s, 'pending')", (customer['customer_id'],))
                mysql.connection.commit()
                cart_id = cursor.lastrowid
            else:
                cart_id = cart['cart_id']

            cursor.execute("SELECT * FROM Cart_Items WHERE cart_id = %s AND product_id = %s", (cart_id, product_id))
            cart_item = cursor.fetchone()

            if cart_item:
                new_quantity = cart_item['quantity'] + 1
                cursor.execute("UPDATE Cart_Items SET quantity = %s WHERE cart_item_id = %s", (new_quantity, cart_item['cart_item_id']))
            else:
                cursor.execute("INSERT INTO Cart_Items (cart_id, product_id, quantity) VALUES (%s, %s, 1)", (cart_id, product_id))

            mysql.connection.commit()
            flash('Product is added to cart!', 'success')
            return redirect(url_for('product_detail', product_id=product_id))

        except Exception as e:
            mysql.connection.rollback()
            flash(f'Error occurred: {str(e)}', 'danger')
            return redirect(url_for('product_detail', product_id=product_id))

        finally:
            if cursor:
                cursor.close()

    try:
        cursor = mysql.connection.cursor(DictCursor)

        cursor.execute("""
            SELECT p.*, c.category_name 
            FROM Products p
            JOIN Categories c ON p.category_id = c.category_id
            WHERE p.product_id = %s
        """, (product_id,))
        product = cursor.fetchone()
        
        if not product:
            flash('Product is not found.', 'danger')
            return redirect(url_for('home'))
        
        cursor.close()
        
        return render_template('product_detail.html', product=product)
    
    except Exception as e:
        flash(f'Error: {e}', 'danger')
        return redirect(url_for('home'))
@app.route('/guest_product_detail/<int:product_id>', methods=['GET', 'POST'])
def guest_product_detail(product_id):
    if request.method == 'POST':
        if 'user_id' not in session or session.get('role') != 'customer':
            flash('You need to login to add items to your cart.', 'danger')
            return redirect(url_for('login'))

        try:
            cursor = mysql.connection.cursor(DictCursor)
            cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
            customer = cursor.fetchone()

            if not customer:
                flash('Customer information is not available.', 'danger')
                return redirect(url_for('customer_dashboard'))

            cursor.execute("SELECT cart_id FROM Cart WHERE customer_id = %s AND cart_status = 'pending'", (customer['customer_id'],))
            cart = cursor.fetchone()

            if not cart:
                cursor.execute("INSERT INTO Cart (customer_id, cart_status) VALUES (%s, 'pending')", (customer['customer_id'],))
                mysql.connection.commit()
                cart_id = cursor.lastrowid
            else:
                cart_id = cart['cart_id']

            cursor.execute("SELECT * FROM Cart_Items WHERE cart_id = %s AND product_id = %s", (cart_id, product_id))
            cart_item = cursor.fetchone()

            if cart_item:
                new_quantity = cart_item['quantity'] + 1
                cursor.execute("UPDATE Cart_Items SET quantity = %s WHERE cart_item_id = %s", (new_quantity, cart_item['cart_item_id']))
            else:
                cursor.execute("INSERT INTO Cart_Items (cart_id, product_id, quantity) VALUES (%s, %s, 1)", (cart_id, product_id))

            mysql.connection.commit()
            flash('Product is added to cart!', 'success')
            return redirect(url_for('guest_product_detail.html', product_id=product_id))

        except Exception as e:
            mysql.connection.rollback()
            flash(f'Error occurred: {str(e)}', 'danger')
            return redirect(url_for('guest_product_detail.html', product_id=product_id))

        finally:
            if cursor:
                cursor.close()

    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT p.*, c.category_name 
            FROM Products p
            JOIN Categories c ON p.category_id = c.category_id
            WHERE p.product_id = %s
        """, (product_id,))
        product = cursor.fetchone()
        
        if not product:
            flash('Product is not found.', 'danger')
            return redirect(url_for('home'))
        
        cursor.close()
        return render_template('guest_product_detail.html', product=product)
    
    except Exception as e:
        flash(f'Error: {e}', 'danger')
        return redirect(url_for('home'))
#------------------------------------------CART-----------------------------------------------
@app.route('/cart') 
def cart():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to see your cart.', 'danger')
        return redirect(url_for('login'))

    customer_id = session['customer_id']
    cursor = None
    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT cart_id 
            FROM Cart 
            WHERE customer_id = %s AND cart_status = 'pending'
        """, (customer_id,))
        cart = cursor.fetchone()

        if not cart:
            return render_template('cart.html', cart_items=None, total_price=0)
      
        cursor.execute("""
            select 
                Products.product_id, 
                Products.product_name, 
                Products.price,
                Products.image_url, 
                Cart_Items.quantity, 
                Cart_Items.cart_item_id 
            from Cart_Items 
            JOIN Products ON Cart_Items.product_id = Products.product_id 
            where cart_id = %s
        """, (cart['cart_id'],))
        cart_items = cursor.fetchall()

        total_price = calculate_total_price(cart_items)

        return render_template('cart.html', cart_items=cart_items, total_price=total_price)

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return render_template('cart.html', cart_items=None, total_price=0)


    finally:
        if cursor:
            cursor.close()

@app.route('/add_to_cart', methods=['POST'])
def add_to_cart():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to add items to your cart.', 'danger')
        return redirect(url_for('login'))

    product_id = request.form.get('product_id')
    if not product_id:
        flash('Product not found.', 'danger')
        return redirect(url_for('customer_dashboard'))

    customer_id = session['user_id']

    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (customer_id,)) 
        customer = cursor.fetchone()
        if not customer:
            flash('Customer information is unavailable.', 'danger')
            return redirect(url_for('customer_dashboard'))

        cursor.execute("SELECT cart_id FROM Cart WHERE customer_id = %s AND cart_status = 'pending'", (customer['customer_id'],))
        cart = cursor.fetchone()

        if not cart:
            cursor.execute("INSERT INTO Cart (customer_id, cart_status) VALUES (%s, 'pending')", (customer['customer_id'],))
            mysql.connection.commit()
            cart_id = cursor.lastrowid
        else:
            cart_id = cart['cart_id']

        cursor.execute("SELECT * FROM Cart_Items WHERE cart_id = %s AND product_id = %s", (cart_id, product_id))
        cart_item = cursor.fetchone()

        if cart_item:
            new_quantity = cart_item['quantity'] + 1
            cursor.execute("UPDATE Cart_Items SET quantity = %s WHERE cart_item_id = %s", (new_quantity, cart_item['cart_item_id']))
        else:
            cursor.execute("INSERT INTO Cart_Items (cart_id, product_id, quantity) VALUES (%s, %s, %s)", (cart_id, product_id, 1))

        mysql.connection.commit()
        flash('Product is added into cart!', 'success')

    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')

    finally:
        if cursor:
            cursor.close()

    return redirect(url_for('customer_dashboard'))


@app.route('/update_cart_item', methods=['POST'])
def update_cart_item():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to update your cart.', 'danger')
        return redirect(url_for('login'))

    cart_item_id = request.form.get('cart_item_id')
    quantity = request.form.get('quantity')
    
    try:
        cursor = mysql.connection.cursor()
        cursor.execute("UPDATE Cart_Items SET quantity = %s where cart_item_id = %s", (quantity, cart_item_id))
        mysql.connection.commit()
        flash('Update your cart successfully!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
    
    return redirect(url_for('cart'))

@app.route('/remove_cart_item', methods=['POST']) 
def delete_cart_item():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to delete items from your cart.', 'danger')
        return redirect(url_for('login'))

    cart_item_id = request.form.get('cart_item_id')
    cursor = None

    try:
        cursor = mysql.connection.cursor(DictCursor)

        cursor.execute("select cart_id from Cart_Items where cart_item_id = %s", (cart_item_id,))
        cart_item = cursor.fetchone()
        if not cart_item:
            flash('Cannot find any items in the cart.', 'danger')
            return redirect(url_for('cart'))

        cart_id = cart_item['cart_id']
        cursor.execute("DELETE from Cart_Items where cart_item_id = %s", (cart_item_id,))
        mysql.connection.commit()
        flash('Delete item from your cart successfully!', 'success')

        cursor.execute("select COUNT(*) AS item_count from Cart_Items where cart_id = %s", (cart_id,))
        result = cursor.fetchone()
        item_count = result['item_count']

        if item_count == 0:
            cursor.execute("DELETE from Cart where cart_id = %s", (cart_id,))
            mysql.connection.commit()
            flash('The cart is empty and is deleted from the system.', 'success')

    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')

    finally:
        if cursor:
            cursor.close()

    return redirect(url_for('cart'))

#------------------------------------------PAYMENT----------------------------------------------
@app.route('/checkout')
def checkout():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login as customer to pay.', 'danger')
        return redirect(url_for('login'))

    user_id = session['user_id']
    cursor = None

    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (user_id,))
        customer = cursor.fetchone()

        if not customer:
            flash('Cannot find any customer information. Please check again.', 'danger')
            return redirect(url_for('customer_dashboard'))

        customer_id = customer['customer_id']
        cursor.execute("SELECT * FROM Cart WHERE customer_id = %s AND cart_status = 'pending'", (customer_id,))
        cart = cursor.fetchone()

        if not cart:
            flash('Cannot find any carts. Please add new items to cart before paying.', 'danger')
            return redirect(url_for('cart'))

        cart_id = cart['cart_id']
        cursor.execute("""
            SELECT 
                Products.product_id,
                Products.product_name,
                Products.price,
                Cart_Items.quantity
            FROM Cart_Items
            JOIN Products ON Cart_Items.product_id = Products.product_id
            WHERE cart_id = %s
        """, (cart_id,))
        cart_items = cursor.fetchall()

        if not cart_items:
            flash('Your cart is currently empty. Please add new items into the cart.', 'danger')
            return redirect(url_for('cart'))

        subtotal = float(sum(float(item['price']) * float(item['quantity']) for item in cart_items))
        shipping_fee = 20000
        total = float(subtotal)

        cursor.execute("""
            SELECT mt.tier_id, mt.is_active, mtr.discount
            FROM MembershipTracking mt
            JOIN Membership_Tiers mtr ON mt.tier_id = mtr.tier_id
            WHERE mt.customer_id = %s AND mt.is_active = 1 
            AND CURDATE() BETWEEN mt.start_date AND mt.end_date
        """, (customer_id,))
        membership = cursor.fetchone()
        
        cursor.execute("""
            SELECT p.*
            FROM Promotions p
            WHERE NOW() BETWEEN p.start_date AND p.end_date
            AND (p.max_uses = 0 OR p.current_uses < p.max_uses)
            AND NOT EXISTS (
                SELECT 1 
                FROM Orders o 
                WHERE o.promotion_id = p.promotion_id
                AND o.customer_id = %s
            )
            AND p.promotion_id NOT IN (
                SELECT promotion_id 
                FROM Orders 
                WHERE customer_id = %s 
                AND status != 'cancelled'
            )
        """, (customer_id, customer_id))
        available_promotions = get_active_promotions(customer_id)
        if not available_promotions:
            print("No promotions found for customer:", customer_id)
            
        membership_discount = 0
        if membership and membership['is_active']:
            membership_discount = float(float(subtotal) * float(membership['discount']) / 100)
            total -= float(membership_discount)

        promotion_discount = 0
        if 'applied_promotion' in session:
            promotion = session['applied_promotion']
            promotion_discount = float(float(subtotal) * float(promotion['discount_percentage']) / 100)
            total -= float(promotion_discount)
            flash(f'The promotion {promotion["code"]} is applied for a discount of {"{:,.0f}".format(promotion_discount)} VNĐ', 'success')

        total += float(shipping_fee)

        order = {
            'order_id': 'Pending',
            'total_price': total,
            'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'payment_method': 'Not yet determined',
            'payment_status': 'pending'
        }
        return render_template(
            'checkout.html',
            order=order,
            order_items=cart_items,
            subtotal=subtotal,
            total=total,
            membership_discount=membership_discount,
            promotion_discount=promotion_discount,
            shipping_fee=shipping_fee,
            available_promotions=available_promotions
        )

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))

    finally:
        if cursor:
            cursor.close()

@app.route('/create_order', methods=['POST'])
def create_order():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to pay.', 'danger')
        return redirect(url_for('login'))

    payment_method = request.form.get('payment_method') 
    cursor = None

    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
        customer = cursor.fetchone()
        if not customer:
            raise Exception("Customer not found")
            return redirect(url_for('customer_dashboard'))

        customer_id = customer['customer_id']
        cursor.execute("""
            SELECT * FROM Cart 
            WHERE customer_id = %s AND cart_status = 'pending'
        """, (customer_id,))
        cart = cursor.fetchone()

        if not cart:
            flash('Cannot find any carts.', 'danger')
            return redirect(url_for('customer_dashboard'))

        cart_id = cart['cart_id']
        cursor.execute("""
            SELECT 
                Products.product_id,
                Products.price,
                Cart_Items.quantity
            FROM Cart_Items
            JOIN Products ON Cart_Items.product_id = Products.product_id
            WHERE cart_id = %s
        """, (cart_id,))
        cart_items = cursor.fetchall()

        if not cart_items:
            flash('Your cart is currently empty.', 'danger')
            return redirect(url_for('cart'))

        subtotal = float(sum(float(item['price']) * float(item['quantity']) for item in cart_items))
        shipping_fee = 20000
        total = float(subtotal)

        cursor.execute("""
            SELECT mt.tier_id, mt.is_active, mtr.discount
            FROM MembershipTracking mt
            JOIN Membership_Tiers mtr ON mt.tier_id = mtr.tier_id
            WHERE mt.customer_id = %s AND mt.is_active = 1 
            AND CURDATE() BETWEEN mt.start_date AND mt.end_date
        """, (customer_id,))
        membership = cursor.fetchone()

        membership_discount = 0
        if membership and membership['is_active']:
            membership_discount = float(float(subtotal) * float(membership['discount']) / 100)
            total -= float(membership_discount)

        promotion_id = None
        promotion_discount = 0
        if 'applied_promotion' in session:
            promotion = session['applied_promotion']
            cursor.execute("SELECT promotion_id FROM Promotions WHERE promotion_code = %s", (promotion['code'],))
            promotion_result = cursor.fetchone()
            if promotion_result:
                promotion_id = promotion_result['promotion_id']
                promotion_discount = float(float(subtotal) * float(promotion['discount_percentage']) / 100)
                total -= float(promotion_discount)
                
                cursor.execute("""
                    UPDATE Promotions 
                    SET current_uses = current_uses + 1 
                    WHERE promotion_id = %s
                """, (promotion_id,))
                mysql.connection.commit()
                flash(f'Promotion {promotion["code"]} applied successfully!', 'success')

        total += float(shipping_fee)

        cursor.execute("""
            INSERT INTO Orders (customer_id, total_price, status, promotion_id)
            VALUES (%s, %s, 'processing', %s)
        """, (customer_id, float(total), promotion_id))
        order_id = cursor.lastrowid

        for item in cart_items:
            cursor.execute("""
                INSERT INTO Order_Items (order_id, product_id, quantity, subtotal)
                VALUES (%s, %s, %s, %s)
            """, (order_id, item['product_id'], item['quantity'], item['price'] * item['quantity']))

        if payment_method == 'cash':
            cursor.execute("""
                INSERT INTO Payments (order_id, payment_method, payment_status, created_at)
                VALUES (%s, 'cash', 'pending', NOW())
            """, (order_id,))
            
        elif payment_method == 'deposit':
            payment_proof = request.files.get('payment_proof')
            if payment_proof:
                filename = secure_filename(f"payment_proof_{order_id}_{int(time.time())}.jpg")
                payment_proof.save(os.path.join('static/uploads/payments', filename))
                
                cursor.execute("""
                    INSERT INTO Payments (order_id, payment_method, payment_status, payment_proof, created_at)
                    VALUES (%s, 'deposit', 'pending', %s, NOW())
                """, (order_id, filename))
            else:
                flash('Please upload payment proof for deposit method', 'danger')
                return redirect(url_for('checkout'))

        mysql.connection.commit()
        flash('Order created successfully!', 'success')

        cursor.execute("DELETE FROM Cart_Items WHERE cart_id = %s", (cart_id,))
        cursor.execute("DELETE FROM Cart WHERE cart_id = %s", (cart_id,))
        if 'applied_promotion' in session:
            session.pop('applied_promotion')
        mysql.connection.commit()

        return redirect(url_for('order_confirmation', order_id=order_id))

    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('checkout'))
    finally:
        if cursor:
            cursor.close()


@app.route('/order_confirmation/<int:order_id>')
def order_confirmation(order_id):
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to see your order confirmation.', 'danger')
        return redirect(url_for('login'))

    cursor = None
    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT 
                o.order_id,
                o.total_price,
                o.created_at,
                o.status,
                COALESCE(p.payment_method, 'pending') as payment_method,
                COALESCE(p.payment_status, 'pending') as payment_status
            FROM Orders o
            LEFT JOIN Payments p ON o.order_id = p.order_id
            WHERE o.order_id = %s AND o.customer_id = (
                SELECT customer_id FROM Customers WHERE user_id = %s
            )
        """, (order_id, session['user_id']))
        
        order = cursor.fetchone()
        
        if not order:
            flash('Cannot find any orders.', 'danger')
            return redirect(url_for('customer_dashboard'))

        cursor.execute("""
            SELECT 
                p.product_name,
                oi.quantity,
                oi.subtotal
            FROM Order_Items oi
            JOIN Products p ON oi.product_id = p.product_id
            WHERE oi.order_id = %s
        """, (order_id,))
        order_items = cursor.fetchall()

        return render_template('order_confirmation.html', order=order, order_items=order_items)

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))

    finally:
        if cursor:
            cursor.close()
#------------------------------------------DELIVERY----------------------------------------------
@app.route('/order_history/<int:customer_id>')
def order_history(customer_id):
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to see your order history.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
    customer = cursor.fetchone()
    
    if not customer or customer['customer_id'] != customer_id:
        flash('You are not authorized to view this order history.', 'danger')
        return redirect(url_for('customer_dashboard'))
    
    try:
        cursor.execute("""
            SELECT 
                o.order_id,
                o.total_price,
                o.status,
                o.created_at,
                p.payment_method,
                p.payment_status,
                cr.reason_text,
                EXISTS(
                    SELECT 1 
                    FROM Reviews r 
                    JOIN Order_Items oi ON r.product_id = oi.product_id 
                    WHERE oi.order_id = o.order_id
                ) AS has_review
            FROM Orders o
            LEFT JOIN Payments p ON o.order_id = p.order_id
            LEFT JOIN CancellationReasons cr ON o.reason_id = cr.reason_id
            WHERE o.customer_id = %s
            ORDER BY o.created_at DESC
        """, (customer_id,))
        orders = cursor.fetchall()
        return render_template('order_history.html', orders=orders)

    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))
    
    finally:
        if cursor:
            cursor.close()

@app.route('/order_detail/<int:customer_id>/<int:order_id>')
def order_detail(customer_id, order_id):
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You are not allowed to access this order.', 'danger')
        return redirect(url_for('login'))
    
    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
        customer = cursor.fetchone()
        
        if not customer:
            flash('Customer information not found.', 'danger')
            return redirect(url_for('customer_dashboard'))

        cursor.execute("""
            SELECT 
                o.order_id,
                o.total_price,
                o.status,
                o.created_at,
                o.promotion_id,
                c.fullname,
                c.phone_number,
                p.payment_method,
                p.payment_status
            FROM Orders o
            JOIN Customers c ON o.customer_id = c.customer_id
            LEFT JOIN Payments p ON o.order_id = p.order_id
            WHERE o.order_id = %s AND o.customer_id = %s
        """, (order_id, customer['customer_id']))
        order = cursor.fetchone()

        if not order:
            flash('Order not found or you are not authorized to view this order.', 'danger')
            return redirect(url_for('customer_dashboard'))

        if order['promotion_id']:
            cursor.execute("SELECT * FROM Promotions WHERE promotion_id = %s", (order['promotion_id'],))
            promotion = cursor.fetchone()
        else:
            promotion = None
        
        cursor.execute("""
            SELECT 
                p.product_name,
                oi.quantity,
                oi.subtotal
            FROM Order_Items oi
            JOIN Products p ON oi.product_id = p.product_id
            WHERE oi.order_id = %s
        """, (order_id,))
        order_items = cursor.fetchall()
        
        cursor.execute("""
            SELECT 1
            FROM Reviews r
            JOIN Order_Items oi ON r.product_id = oi.product_id
            WHERE oi.order_id = %s AND r.customer_id
        """, (order_id,))
        has_review = cursor.fetchone()

        cursor.execute("""
                SELECT 
                    r.review_id,
                    r.product_id,
                    p.product_name,
                    r.rating,
                    r.comment,
                    r.created_at
                FROM Reviews r
                JOIN Products p ON r.product_id = p.product_id
                JOIN Order_Items oi ON p.product_id = oi.product_id
                WHERE oi.order_id = %s AND r.customer_id = %s
            """, (order['order_id'], customer_id))
        order['reviews'] = cursor.fetchall()
        order['has_review'] = len(order['reviews']) > 0
        return render_template('order_detail.html', 
                             order=order, 
                             order_items=order_items, promotion=promotion)
    
    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('order_history', customer_id=customer['customer_id']))
    
    finally:
        if cursor:
            cursor.close()
#------------------------------------------PROFILE----------------------------------------------
@app.route ('/customer_profile', methods=['GET'])
def customer_profile():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to check your profile.', 'danger')
        return redirect(url_for('login'))

    user_id = session['user_id']
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute(""" select c.*, u.email 
                   from Customers c
                   join Users u on c.user_id = u.user_id
                   where c.user_id = %s
                   """, (user_id,))
    customer = cursor.fetchone()
    return render_template('customer_profile.html', customer=customer)

@app.route('/update_profile', methods=['GET', 'POST'])
def update_profile():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to update your profile.', 'danger')
        return redirect(url_for('login'))

    user_id = session['user_id']
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("SELECT * FROM Customers WHERE user_id = %s", (user_id,))
        customer = cursor.fetchone()

        if request.method == 'POST':
            fullname = request.form.get('fullname')
            phone_number = request.form.get('phone_number')
            address = request.form.get('address')

            if not fullname or not phone_number or not address:
                flash('Please enter all fields.', 'danger')
            else:
                try:
                    cursor.execute(
                        "UPDATE Customers SET fullname = %s, phone_number = %s, address = %s WHERE user_id = %s",
                        (fullname, phone_number, address, user_id)
                    )
                    mysql.connection.commit()
                    flash('Update your information successfully!', 'success')
                    cursor.execute("SELECT * FROM Customers WHERE user_id = %s", (user_id,))
                    customer = cursor.fetchone()
                except Exception as e:
                    mysql.connection.rollback()
                    flash(f'Error occurred: {str(e)}', 'danger')
        cursor.execute("""
            SELECT u.email FROM Users u
            WHERE u.user_id = %s
        """, (user_id,))
        user = cursor.fetchone()
        if user and customer:
            customer['email'] = user['email']

        return render_template('update_profile.html', customer=customer)
    finally:
        cursor.close()

@app.route('/register_membership', methods=['GET', 'POST'])
def register_membership():
    try:
        if 'user_id' not in session or session.get('role') != 'customer':
            flash('You need to login to register memberships.', 'danger')
            return redirect(url_for('login'))

        user_id = session['user_id']
        cursor = mysql.connection.cursor(DictCursor)

        cursor.execute("""
          SELECT cm.*, mt.tier_name
          FROM MembershipTracking cm
          JOIN Membership_Tiers mt ON cm.tier_id = mt.tier_id
          WHERE cm.customer_id = (SELECT customer_id FROM Customers WHERE user_id = %s) AND cm.is_active = 1
        """, (user_id,))
        existing_membership = cursor.fetchone()

        if existing_membership:
            flash(f"You are already a {existing_membership['tier_name']} member. Cannot register again.", 'info')
            return redirect(url_for('customer_dashboard'))

        cursor.execute("SELECT * FROM Membership_Tiers")
        membership_tiers = cursor.fetchall()

        if request.method == 'POST':
            tier_id = 1 
            cursor.execute("SELECT customer_id from Customers where user_id = %s", (user_id,))
            customer = cursor.fetchone()

            if not customer:
                flash('Cannot find any customer information.', 'danger')
                return redirect(url_for('customer_dashboard'))

            customer_id = customer['customer_id']

            cursor.execute("SELECT * FROM Membership_Tiers WHERE tier_id = %s", (tier_id,))
            tier = cursor.fetchone()

            if not tier:
                flash(f'Cannot find any member level of {tier_id}.', 'danger')
                return redirect(url_for('register_membership'))
            
            cursor.execute("SELECT * FROM MembershipTracking WHERE customer_id = %s AND is_active = 0", (customer_id,))
            existing_pending_membership = cursor.fetchone()
            if existing_pending_membership:
                flash(f'You had registered to be a member and is waiting for the admin to approve.', 'danger')
                return redirect(url_for('register_membership'))

            if not existing_pending_membership:
                cursor.execute("""
                INSERT INTO MembershipTracking (customer_id, tier_id, start_date, end_date, is_active, total_spent)
                VALUES (%s, %s, CURDATE(), DATE_ADD(CURDATE(), INTERVAL 1 YEAR), 0, 0) 
                """, (customer_id, tier_id))
            mysql.connection.commit()

            flash('Membership registration successfully. Please wait for the admin to approve.', 'success')
            return redirect(url_for('customer_dashboard'))

        return render_template('register_membership.html', membership_tiers=membership_tiers)

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))

    finally:
        if 'cursor' in locals():
            cursor.close()

#------------------------------------------WISHLIST--------------------------------------------------------
@app.route('/wishlist', methods=['GET', 'POST'])
def wishlist():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to see your wishlist.', 'danger')
        return redirect(url_for('login'))

    user_id = session['user_id']
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("""
            SELECT p.product_id, p.product_name, p.price, p.image_url
            FROM Wishlists w
            JOIN Products p ON w.product_id = p.product_id
            WHERE w.customer_id = (SELECT customer_id FROM Customers WHERE user_id = %s)
        """, (user_id,))
        wishlist = cursor.fetchall()
        return render_template('wishlist.html', wishlist=wishlist)
    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('customer_dashboard'))
    finally:
        cursor.close()
    return render_template('wishlist.html', wishlist=wishlist)
@app.route('/add_to_wishlist', methods=['POST'])
def add_to_wishlist():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to add items to your wishlist.', 'danger')
        return redirect(url_for('login'))

    product_id = request.form.get('product_id')
    if not product_id:
        flash('Cannot find any items.', 'danger')
        return redirect(url_for('customer_dashboard'))

    user_id = session['user_id']
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("select * from Wishlists where customer_id = (select customer_id from Customers where user_id = %s) and product_id = %s", (user_id, product_id))
        product = cursor.fetchone()
        if product:
            flash('Product is already in the wishlist.', 'info')
        else:
            cursor.execute("INSERT INTO Wishlists (customer_id, product_id) VALUES ((SELECT customer_id FROM Customers WHERE user_id = %s), %s)", (user_id, product_id))
            mysql.connection.commit()
            flash('Product is added into the wishlist!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
    finally:
        cursor.close()
    return redirect(url_for('customer_dashboard'))
@app.route('/remove_from_wishlist', methods=['POST'])
def remove_from_wishlist():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to manage your wishlist.', 'danger')
        return redirect(url_for('login'))
    
    product_id = request.form.get('product_id')
    if not product_id:
        flash('Product not found.', 'danger')
        return redirect(url_for('wishlist'))
    
    user_id = session['user_id']
    cursor = mysql.connection.cursor()
    try:
        cursor.execute("DELETE FROM Wishlist WHERE user_id = %s AND product_id = %s", (user_id, product_id))
        mysql.connection.commit()
        flash('Product removed from your wishlist.', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
    finally:
        cursor.close()
    
    return redirect(url_for('wishlist'))
#-------------------------------------------REVIEWS----------------------------------------------------------
@app.route('/reviews')
def reviews():
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("""
        SELECT r.*, p.product_name, c.fullname as customer_name
        FROM Reviews r
        JOIN Products p ON r.product_id = p.product_id
        JOIN Customers c ON r.customer_id = c.customer_id
        ORDER BY r.created_at DESC
    """)
    reviews = cursor.fetchall()
    cursor.close()
    return render_template('reviews.html', reviews=reviews)

@app.route('/add_review/<int:customer_id>/<int:order_id>', methods=['GET', 'POST'])
def add_review(customer_id, order_id):
    if 'user_id' not in session or session.get('role')!= 'customer':
        flash('You need to login to add reviews.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("""
            SELECT o.order_id, oi.product_id, p.product_name
            FROM Orders o
            JOIN Order_Items oi ON o.order_id = oi.order_id
            JOIN Products p ON oi.product_id = p.product_id
            WHERE o.order_id = %s 
            AND o.customer_id = (SELECT customer_id FROM Customers WHERE user_id = %s)
            AND o.status = 'completed'
        """, (order_id, session['user_id']))
        
        order_items = cursor.fetchall()
        
        if not order_items:
            flash('You can only review products from your completed orders.', 'danger')
            return redirect(url_for('order_history', customer_id=session['customer_id']))

        if request.method == 'POST':
            has_rating = False
            for item in order_items:
                if request.form.get(f'rating_{item['product_id']}'):
                    has_rating = True
                    break
            
            if not has_rating:
                flash('Please provide at least one rating.', 'danger')
                return render_template('add_review.html', order_items=order_items, order_id=order_id)

            for item in order_items:
                product_id = item['product_id']
                rating = request.form.get(f'rating_{product_id}')
                comment = request.form.get(f'comment_{product_id}')
                
                if rating: 
                    cursor.execute("""
                        INSERT INTO Reviews (customer_id, product_id, rating, comment, created_at)
                        VALUES (
                            (SELECT customer_id FROM Customers WHERE user_id = %s),
                            %s, %s, %s, NOW()
                        )
                    """, (session['user_id'], product_id, rating, comment))
            
            mysql.connection.commit()
            flash('Reviews added successfully!', 'success')
            return redirect(url_for('order_history', customer_id=session['customer_id']))

        return render_template('add_review.html', order_items=order_items, order_id=order_id)

    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('order_history', customer_id=session['customer_id'], order_id=order_id))
    finally:
        if cursor:
            cursor.close()

#-------------------------------------------OTHERS----------------------------------------------------------
@app.route('/about_us')
def about_us():
    return render_template('about_us.html')
#-------------------------------------------ADMIN: PRODUCTS------------------------------------------------
@app.route('/admin_products', methods=['GET', 'POST'])
def admin_products():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("SELECT * FROM Products")
    products = cursor.fetchall()
    cursor.close()
    
    return render_template('admin/admin_products.html', products=products)

@app.route('/add_product', methods=['GET', 'POST'])
def add_product():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("SELECT * FROM Categories")
    categories = cursor.fetchall()
    
    if request.method == 'POST':
        product_name = request.form.get('product_name')
        price = request.form.get('price')
        category_id = request.form.get('category_id')
        image_file = request.files.get('image_file')
        
        if not product_name or not price or not category_id or not image_file:
            flash('Please enter all fields and upload an image.', 'danger')
        elif not image_file.filename:
            flash('No selected file', 'danger')
        else:
            try:
                upload_dir = os.path.join('static', 'images')
                os.makedirs(upload_dir, exist_ok=True)
                
                filename = secure_filename(f"{product_name}_{int(time.time())}.{image_file.filename.split('.')[-1]}")
                image_path = os.path.join(upload_dir, filename)
                image_file.save(image_path)

                image_url = f"images/{filename}"
                
                cursor.execute("INSERT INTO Products (product_name, price, image_url, category_id) VALUES (%s, %s, %s, %s)",
                               (product_name, price, image_url, category_id))
                mysql.connection.commit()
                flash('Product added successfully!', 'success')
                return redirect(url_for('admin_products'))
            except Exception as e:
                mysql.connection.rollback()
                flash(f'Error occurred: {str(e)}', 'danger')
            finally:
                cursor.close()
    return render_template('admin/add_product.html', categories=categories)

@app.route('/edit_product/<int:product_id>', methods=['GET', 'POST'])
def edit_product(product_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("SELECT * FROM Products WHERE product_id = %s", (product_id,))
    product = cursor.fetchone()
    
    if request.method == 'POST':
        product_name = request.form.get('product_name')
        price = request.form.get('price')
        category_id = request.form.get('category_id')
        image_file = request.files.get('image_file')
        
        if not product_name or not price or not category_id:
            flash('Please enter all required fields.', 'danger')
        else:
            try:
                if image_file and image_file.filename:
                    upload_dir = os.path.join('static', 'images')
                    os.makedirs(upload_dir, exist_ok=True)
                    
                    filename = secure_filename(f"{product_name}_{int(time.time())}.{image_file.filename.split('.')[-1]}")
                    image_path = os.path.join(upload_dir, filename)
                    image_file.save(image_path)
                    image_url = f"images/{filename}"
                else:
                    image_url = product['image_url']
                
                cursor.execute("""
                    UPDATE Products 
                    SET product_name = %s, 
                        price = %s, 
                        image_url = %s, 
                        category_id = %s 
                    WHERE product_id = %s
                """, (product_name, price, image_url, category_id, product_id))
                
                mysql.connection.commit()
                flash('Product updated successfully!', 'success')
                return redirect(url_for('admin_products'))
                
            except Exception as e:
                mysql.connection.rollback()
                flash(f'Error occurred: {str(e)}', 'danger')
    else:
        cursor.execute("SELECT * FROM Categories")
        categories = cursor.fetchall()
        cursor.close()
    
    return render_template('admin/edit_product.html', 
                         product=product, 
                         categories=categories)

@app.route('/delete_product/<int:product_id>', methods=['POST'])
def delete_product(product_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("DELETE FROM Products WHERE product_id = %s", (product_id,))
        mysql.connection.commit()
        flash('Deleting a product successfully!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
    finally:
        cursor.close()
    
    return redirect(url_for('admin_products'))

#-------------------------------------------ADMIN: CATEGORIES------------------------------------------------
@app.route('/admin_categories', methods=['GET', 'POST'])
def admin_categories():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("select * from Categories")
    categories = cursor.fetchall()
    return render_template('admin/admin_categories.html', categories=categories)

@app.route('/add_category', methods=['GET', 'POST'])
def add_category():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    if request.method == 'POST':
        category_name = request.form.get('category_name')
        description = request.form.get('description')

        if not category_name:
            flash('Please enter all fields.', 'danger')
        else:
            try:
                cursor.execute("INSERT INTO Categories (category_name, description) VALUES (%s, %s)", 
                (category_name, description))
                mysql.connection.commit()
                flash('Category added successfully!', 'success')
                return redirect(url_for('admin_categories'))
            except Exception as e:
                mysql.connection.rollback()
                flash(f'Error occurred: {str(e)}', 'danger')
            finally:
                cursor.close()
    
    return render_template('admin/add_category.html')

@app.route('/edit_category/<int:category_id>', methods=['GET', 'POST'])
def edit_category(category_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("SELECT * FROM Categories WHERE category_id = %s", (category_id,))
    category = cursor.fetchone()

    if request.method == 'POST':
        category_name = request.form.get('category_name')
        description = request.form.get('description')

        if not category_name:
            flash('Please enter all fields.', 'danger')
        else:
            try:
                cursor.execute("UPDATE Categories SET category_name = %s, description = %s WHERE category_id = %s", 
                             (category_name, description, category_id))
                mysql.connection.commit()
                flash('Update an existing category successfully!', 'success')
                return redirect(url_for('admin_categories'))
            except Exception as e:
                mysql.connection.rollback()
                flash(f'Error occurred: {str(e)}', 'danger')
            finally:
                if cursor:
                    cursor.close() 
    return render_template('admin/edit_category.html', category=category)

@app.route('/delete_category/<int:category_id>', methods=['GET', 'POST'])
def delete_category(category_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor()
    try:
        cursor.execute("SELECT * FROM Products WHERE category_id = %s", (category_id,))
        products = cursor.fetchall()
        if products:
            flash('Cannot delete the category because there are still products associated with it.', 'danger')
            return redirect(url_for('admin_categories'))
        
        cursor.execute("DELETE FROM Categories WHERE category_id = %s", (category_id,))
        mysql.connection.commit()
        flash('Category deleted successfully!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
    finally:
        cursor.close()
    
    return redirect(url_for('admin_categories'))

#-------------------------------------------ADMIN: ORDERS------------------------------------------------
@app.route('/confirm_order', methods=['GET', 'POST'])
def confirm_order():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = None
    try:
        cursor = mysql.connection.cursor()
        if request.method == 'POST':
            order_id = request.form.get('order_id')

            if not order_id:
                flash('Cannot find any orders.', 'danger')
                return redirect(url_for('admin_dashboard'))

            cursor.execute("""
                UPDATE Orders o
                JOIN Payments p ON o.order_id = p.order_id
                SET o.status = 'accepted', p.payment_status = CASE
                    WHEN p.payment_method = 'deposit' THEN 'paid'
                    ELSE p.payment_status
                    END
                WHERE order_id = %s AND status = 'processing'
            """, (order_id,))
            mysql.connection.commit()
            flash(f'Order #{order_id} is confirmed.', 'success')
            return redirect(url_for('confirm_order'))

        cursor.execute("""
            SELECT 
                Orders.order_id,
                Orders.total_price,
                Orders.created_at,
                Orders.status,
                Customers.fullname,
                Customers.phone_number,
                Payments.payment_method
            FROM Orders
            JOIN Customers ON Orders.customer_id = Customers.customer_id
            JOIN Payments ON Orders.order_id = Payments.order_id
            WHERE Orders.status = 'processing'
            ORDER BY Orders.created_at DESC
        """)
        processing_orders = cursor.fetchall()
        return render_template('admin/confirm_order.html', orders=processing_orders)

    except Exception as e:
        if cursor:
            mysql.connection.rollback()
        flash(f'Error occurred when confirming the order: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))
    finally:
        if cursor:
            cursor.close()

@app.route('/confirm_delivery', methods=['GET', 'POST'])
def confirm_delivery():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = None
    try:
        cursor = mysql.connection.cursor(DictCursor)
        if request.method == 'POST':
            order_id = request.form.get('order_id')
            if not order_id:
                flash('Cannot find any orders.', 'danger')
                return redirect(url_for('confirm_delivery'))

            cursor.execute("""
                SELECT o.*, c.fullname, c.phone_number
                FROM Orders o
                JOIN Customers c ON o.customer_id = c.customer_id
                WHERE order_id = %s AND o.status = 'accepted'
            """, (order_id,))
            order = cursor.fetchone()

            if not order:
                mysql.connection.rollback()
                flash('Order has been processed or delivered. Please choose another order.', 'danger')
                return redirect(url_for('confirm_delivery'))

            cursor.execute("""
                UPDATE Orders
                SET status = 'delivering'
                WHERE order_id = %s
            """, (order_id,))
            mysql.connection.commit()

            flash(f'Order #{order_id} of {order["fullname"]} is ready to be delivered and completed.', 'success')
            return redirect(url_for('confirm_delivery'))

        cursor.execute("""
            SELECT
                o.order_id,
                o.created_at,
                o.total_price,
                o.status,
                c.fullname,
                c.phone_number,
                c.address,
                p.payment_method
            FROM Orders o
            JOIN Customers c ON o.customer_id = c.customer_id
            JOIN Payments p ON o.order_id = p.order_id
            WHERE o.status = 'accepted'
            ORDER BY o.created_at ASC
        """)
        accepted_orders = cursor.fetchall()
        return render_template('admin/confirm_delivery.html', orders=accepted_orders)

    except Exception as e:
        if cursor:
            mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

    finally:
        if cursor:
            cursor.close()    

@app.route('/admin_orders', methods=['GET', 'POST'])
def admin_orders():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = None
    try:
        cursor = mysql.connection.cursor(DictCursor)
        if request.method == 'POST':
            order_id = request.form.get('order_id')

            if not order_id:
                flash('Cannot find any đơn hàng.', 'danger')
                return redirect(url_for('admin_dashboard'))

            cursor.execute("""
                SELECT o.*, p.payment_method, p.payment_status
                FROM Orders o
                JOIN Payments p ON o.order_id = p.order_id
                WHERE o.order_id = %s
            """, (order_id,))
            order = cursor.fetchone()
            
            if not order:
                flash('Cannot find any orders.', 'danger')
                return redirect(url_for('admin_orders'))

            if order['status'] == 'processing':
                cursor.execute("""
                    UPDATE Orders o
                    JOIN Payments p ON o.order_id = p.order_id
                    SET o.status = 'accepted',
                        p.payment_status = CASE
                            WHEN p.payment_method = 'deposit' THEN 'paid'
                            ELSE p.payment_status
                        END
                    WHERE o.order_id = %s
                """, (order_id,))
            elif order['status'] == 'accepted':
                cursor.execute("UPDATE Orders SET status = 'delivering' WHERE order_id = %s", (order_id,))
            elif order['status'] == 'delivering':
                cursor.execute("""
                    UPDATE Orders o
                    JOIN Payments p ON o.order_id = p.order_id 
                    SET o.status = 'completed',
                        p.payment_status = CASE
                            WHEN p.payment_method = 'cash' THEN 'paid'
                            ELSE p.payment_status 
                        END
                    WHERE o.order_id = %s 
                """, (order_id,))

            mysql.connection.commit()
            flash('Successfully updated the order status!', 'success')
            return redirect(url_for('admin_orders'))

        cursor.execute("""
            SELECT 
                o.order_id,
                o.total_price,
                o.created_at,
                o.status,
                c.fullname,
                c.phone_number,
                c.address,
                p.payment_method,
                p.payment_proof
            FROM Orders o
            JOIN Customers c ON o.customer_id = c.customer_id
            LEFT JOIN Payments p ON o.order_id = p.order_id
            ORDER BY 
                CASE 
                    WHEN o.status = 'processing' THEN 1
                    WHEN o.status = 'accepted' THEN 2
                    WHEN o.status = 'delivering' THEN 3
                    WHEN o.status = 'completed' THEN 4
                    ELSE 5
                END,
                o.created_at DESC
        """)
        orders = cursor.fetchall()
        
        cursor.execute("SELECT * FROM CancellationReasons")
        reasons = cursor.fetchall()
        
        return render_template('admin/admin_orders.html', orders=orders, reasons=reasons)

    except Exception as e:
        if cursor:
            mysql.connection.rollback()
        flash(f'Error occurred khi confirming the order: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))
    finally:
        if cursor:
            cursor.close()

@app.route('/admin_order_detail/<int:order_id>', methods=['GET', 'POST'])
def admin_order_detail (order_id):
    if 'user_id' not in session or session.get('role')!= 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT 
                o.order_id,
                o.total_price,
                o.status,
                o.created_at,
                c.fullname,
                c.phone_number,
                p.payment_method,
                p.payment_status,
                p.payment_proof
            FROM Orders o
            JOIN Customers c ON o.customer_id = c.customer_id
            LEFT JOIN Payments p ON o.order_id = p.order_id
            WHERE o.order_id = %s
        """, (order_id,))
        order = cursor.fetchone()

        if not order:
            flash('Order not found.', 'danger')
            return redirect(url_for('admin_orders'))
        cursor.execute("""
            SELECT 
                p.product_name,
                oi.quantity,
                oi.subtotal
            FROM Order_Items oi
            JOIN Products p ON oi.product_id = p.product_id
            WHERE oi.order_id = %s
        """, (order_id,))
        order_items = cursor.fetchall()

        return render_template('admin/admin_order_detail.html', 
                             order=order, 
                             order_items=order_items)
    
    except Exception as e:
        flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('admin_orders'))
    
    finally:
        if cursor:
            cursor.close()

@app.route('/cancel_order', methods=['GET', 'POST'])
def cancel_order():
    if 'user_id' not in session or session.get('role')!= 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = None
    try:
        cursor = mysql.connection.cursor(DictCursor)
        if request.method == 'POST':
            order_id = request.form.get('order_id')
            reason_id = request.form.get('reason_id')

            if not order_id:
                flash('Cannot find any orders.', 'danger')
                return redirect(url_for('admin_orders'))

            if not reason_id:
                flash('Please select a cancellation reason.', 'danger')
                return redirect(url_for('admin_orders'))

            cursor.execute("""
                SELECT o.*, p.payment_method, p.payment_status
                FROM Orders o
                JOIN Payments p ON o.order_id = p.order_id
                WHERE o.order_id = %s AND p.payment_method = 'deposit'
            """, (order_id,))
            order = cursor.fetchone()
            
            if not order:
                flash('Cannot find any orders or the payment method of that order is not bank transfer.', 'danger')
                return redirect(url_for('admin_orders'))

            cursor.execute("""
                UPDATE Orders o
                JOIN Payments p ON o.order_id = p.order_id
                SET o.status = 'cancelled',
                    p.payment_status = 'cancelled',
                    o.reason_id = %s
                WHERE o.order_id = %s
            """, (reason_id, order_id,))

            mysql.connection.commit()
            flash('Order is cancelled successfully.', 'success')
            return redirect(url_for('admin_orders'))

        cursor.execute("SELECT * FROM CancellationReasons")
        reasons = cursor.fetchall()
        return render_template('admin/admin_orders.html', reasons=reasons)

    except Exception as e:
        if cursor:
            mysql.connection.rollback()
        flash(f'Error occurred when cancelling the order: {str(e)}', 'danger')
        return redirect(url_for('admin_orders'))
    finally:
        if cursor:
            cursor.close()     

#-------------------------------------------ADMIN: CUSTOMERS------------------------------------------------
@app.route ('/admin_customers', methods=['GET', 'POST'])
def admin_customers():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute(""" 
            SELECT
                c.customer_id,
                c.fullname,
                c.address,
                c.phone_number,
                c.membership_status,
                COALESCE(mt.tier_name, 'Không có') AS tier_name
            FROM Customers c
            LEFT JOIN MembershipTracking mtr ON c.customer_id = mtr.customer_id AND mtr.is_active = 1
            LEFT JOIN Membership_Tiers mt ON mtr.tier_id = mt.tier_id
        """)
        customers = cursor.fetchall()
        cursor.close()
        return render_template('admin/admin_customers.html', customers=customers)
    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))
#-------------------------------------------ADMIN: MEMBERSHIP------------------------------------------------
@app.route('/admin_membership', methods=['GET', 'POST'])
def admin_membership():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    cursor.execute("""
        SELECT 
            mt.customer_id,
            c.fullname,
            c.phone_number,
            c.address,
            mt.created_at,
            mt.end_date,
            t.tier_name
        FROM MembershipTracking mt
        JOIN Customers c ON mt.customer_id = c.customer_id
        JOIN Membership_Tiers t ON mt.tier_id = t.tier_id
        WHERE mt.is_active = 0
    """)
    membership_requests = cursor.fetchall()
    cursor.close()

    if not membership_requests:
        flash('No new membership request to process.', 'info')

    return render_template('admin/admin_membership.html', membership_requests=membership_requests)

@app.route('/approve_membership/<int:customer_id>', methods=['POST'])
def approve_membership(customer_id):
    try:
        if 'user_id' not in session or session.get('role') != 'admin':
            flash('You need to login as administrator.', 'danger')
            return redirect(url_for('login'))

        cursor = mysql.connection.cursor(DictCursor)

        cursor.execute("""
            SELECT mt.*, t.tier_name, c.fullname 
            FROM MembershipTracking mt
            JOIN Membership_Tiers t ON mt.tier_id = t.tier_id
            JOIN Customers c ON mt.customer_id = c.customer_id
            WHERE mt.customer_id = %s AND mt.is_active = 0
        """, (customer_id,))
        pending_membership = cursor.fetchone()

        if not pending_membership:
            flash('Cannot find any pending membership requests.', 'danger')
            return redirect(url_for('admin_membership'))

        cursor.execute("""
            UPDATE MembershipTracking 
            SET is_active = 1, 
                start_date = CURDATE(), 
                end_date = DATE_ADD(CURDATE(), INTERVAL 1 YEAR)
            WHERE customer_id = %s AND is_active = 0
        """, (customer_id,))

        cursor.execute("UPDATE Customers SET membership_status = 'active' WHERE customer_id = %s", (customer_id,))
        mysql.connection.commit()

        flash(f'Accepted {pending_membership["tier_name"]} membership for {pending_membership["fullname"]}.', 'success')
        return redirect(url_for('admin_membership'))

    except Exception as e:
        if 'cursor' in locals():
            mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_membership'))

    finally:
        if 'cursor' in locals():
            cursor.close()

@app.route('/reject_membership/<int:customer_id>', methods=['POST'])
def reject_membership(customer_id):
    try:
        if 'user_id' not in session or session.get('role') != 'admin':
            flash('You need to login as administrator.', 'danger')
            return redirect(url_for('login'))

        cursor = mysql.connection.cursor(DictCursor)

        cursor.execute("""
            SELECT mt.*, t.tier_name, c.fullname 
            FROM MembershipTracking mt
            JOIN Membership_Tiers t ON mt.tier_id = t.tier_id
            JOIN Customers c ON mt.customer_id = c.customer_id
            WHERE mt.customer_id = %s AND mt.is_active = 0
        """, (customer_id,))
        pending_membership = cursor.fetchone()

        if not pending_membership:
            flash('Cannot find any pending membership requests.', 'danger')
            return redirect(url_for('admin_membership'))

        cursor.execute("DELETE FROM MembershipTracking WHERE customer_id = %s AND is_active = 0", (customer_id,))
        mysql.connection.commit()

        flash(f'Rejected request of {pending_membership["tier_name"]} membership for {pending_membership["fullname"]}.', 'success')
        return redirect(url_for('admin_membership'))

    except Exception as e:
        if 'cursor' in locals():
            mysql.connection.rollback()
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_membership'))

    finally:
        if 'cursor' in locals():
            cursor.close()
            
#-------------------------------------------ADMIN: PROMOTIONS------------------------------------------------
@app.route('/admin_promotions', methods=['GET', 'POST'])
def admin_promotions():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    try:
        if request.method == 'POST':
            return redirect(url_for('admin_promotions'))
        cursor.execute("""
            SELECT *,
                CASE 
                    WHEN NOW() < start_date THEN 'Sắp diễn ra'
                    WHEN NOW() BETWEEN start_date AND end_date THEN 'Đang diễn ra'
                    ELSE 'Đã kết thúc'
                END as status
            FROM Promotions
        """)
        promotions = cursor.fetchall()
        return render_template('admin/admin_promotions.html', promotions=promotions)
    finally:
        cursor.close()

@app.route('/add_promotion', methods=['GET', 'POST'])
def add_promotion():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    try:
        if request.method == 'POST':
            promotion_code = request.form.get('promotion_code', '').strip()
            promotion_name = request.form.get('promotion_name', '').strip()
            discount_percentage = request.form.get('discount_percentage', '').strip()
            start_date = request.form.get('start_date', '').strip()
            end_date = request.form.get('end_date', '').strip()
            max_uses = request.form.get('max_uses', '0').strip()

            if not all([promotion_code, promotion_name, discount_percentage, start_date, end_date]):
                flash('Please fill all required fields', 'danger')
                return render_template('admin/add_promotion.html',
                                    promotion_code=promotion_code,
                                    promotion_name=promotion_name,
                                    discount_percentage=discount_percentage,
                                    start_date=start_date,
                                    end_date=end_date,
                                    max_uses=max_uses)

            try:
                discount_percentage = float(discount_percentage)
                if discount_percentage <= 0 or discount_percentage > 100:
                    raise ValueError
            except ValueError:
                flash('Discount percentage must be a number between 0 and 100', 'danger')
                return render_template('admin/add_promotion.html',
                                    promotion_code=promotion_code,
                                    promotion_name=promotion_name,
                                    discount_percentage=discount_percentage,
                                    start_date=start_date,
                                    end_date=end_date,
                                    max_uses=max_uses)

            if start_date >= end_date:
                flash('End date must be after start date', 'danger')
                return render_template('admin/add_promotion.html',
                                    promotion_code=promotion_code,
                                    promotion_name=promotion_name,
                                    discount_percentage=discount_percentage,
                                    start_date=start_date,
                                    end_date=end_date,
                                    max_uses=max_uses)

            try:
                cursor.execute("""
                    INSERT INTO Promotions (
                        promotion_code, promotion_name, discount_percentage,
                        start_date, end_date, max_uses, current_uses, created_at
                    ) VALUES (%s, %s, %s, %s, %s, %s, 0, NOW())
                """, (promotion_code, promotion_name, discount_percentage, 
                     start_date, end_date, max_uses))
                mysql.connection.commit()
                flash('Promotion added successfully!', 'success')
                return redirect(url_for('admin_promotions'))
            except Exception as e:
                mysql.connection.rollback()
                flash(f'Error adding promotion: {str(e)}', 'danger')
                return render_template('admin/add_promotion.html',
                                    promotion_code=promotion_code,
                                    promotion_name=promotion_name,
                                    discount_percentage=discount_percentage,
                                    start_date=start_date,
                                    end_date=end_date,
                                    max_uses=max_uses)

        return render_template('admin/add_promotion.html')

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_promotions'))
    finally:
        cursor.close()

@app.route('/edit_promotion/<int:promotion_id>', methods=['GET', 'POST'])
def edit_promotion(promotion_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor(DictCursor)
    try:
        if request.method == 'POST':
            promotion_name = request.form.get('promotion_name', '').strip()
            discount_percentage = request.form.get('discount_percentage', '').strip()
            start_date = request.form.get('start_date', '').strip()
            end_date = request.form.get('end_date', '').strip()
            max_uses = request.form.get('max_uses', '0').strip()

            if not all([promotion_name, discount_percentage, start_date, end_date]):
                flash('Please fill all required fields', 'danger')
                return render_template('admin/edit_promotion.html', promotion=promotion)

            try:
                discount_percentage = float(discount_percentage)
                if discount_percentage <= 0 or discount_percentage > 100:
                    raise ValueError
            except ValueError:
                flash('Discount percentage must be a number between 0 and 100', 'danger')
                return render_template('admin/edit_promotion.html', promotion=promotion)

            if start_date >= end_date:
                flash('End date must be after start date', 'danger')
                return render_template('admin/edit_promotion.html', promotion=promotion)

            cursor.execute("""
                UPDATE Promotions 
                SET promotion_name = %s, discount_percentage = %s,
                    start_date = %s, end_date = %s, max_uses = %s
                WHERE promotion_id = %s
            """, (promotion_name, discount_percentage, start_date, end_date, max_uses, promotion_id))
            mysql.connection.commit()
            flash('Updated the promo code successfully!', 'success')
            return redirect(url_for('admin_promotions'))

        cursor.execute("SELECT * FROM Promotions WHERE promotion_id = %s", (promotion_id,))
        promotion = cursor.fetchone()
        if not promotion:
            flash('Cannot find any promotion codes.', 'danger')
            return redirect(url_for('admin_promotions'))
        return render_template('admin/edit_promotion.html', promotion=promotion)
    finally:
        cursor.close()

@app.route('/delete_promotion/<int:promotion_id>', methods=['POST'])
def delete_promotion(promotion_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    cursor = mysql.connection.cursor()
    try:
        cursor.execute("DELETE FROM Promotions WHERE promotion_id = %s", (promotion_id,))
        mysql.connection.commit()
        flash('Delete that promotion code successfully!', 'success')
    except Exception as e:
        mysql.connection.rollback()
        flash(f'Lỗi: {str(e)}', 'danger')
    finally:
        cursor.close()
    return redirect(url_for('admin_promotions'))

#-------------------------------------------CUSTOMER: PROMOTIONS------------------------------------------------
@app.route('/get_available_promotions')
def get_available_promotions():
    if 'user_id' not in session or session.get('role') != 'customer':
        return jsonify({'error': 'Unauthorized'}), 401

    try:
        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT
                p.promotion_id,
                p.promotion_code,
                p.promotion_name,
                p.discount_percentage,
                p.start_date,
                p.end_date,
                p.max_uses,
                p.current_uses
            FROM Promotions p
            WHERE CURDATE() BETWEEN p.start_date AND p.end_date
            AND (p.max_uses = 0 OR p.current_uses < p.max_uses)
            AND NOT EXISTS (
                SELECT 1 
                FROM Orders o 
                WHERE o.promotion_id = p.promotion_id
                AND o.customer_id = %s
                AND o.status != 'cancelled'
            )
            ORDER BY p.discount_percentage DESC
        """, (session['user_id'],))
        promotions = cursor.fetchall()
        return jsonify(promotions)
    except Exception as e:
        print(f"Error fetching promotions: {str(e)}")
        return jsonify({'error': str(e)}), 500
    finally:
        if cursor:
            cursor.close()

@app.route('/apply_promotion_code', methods=['POST'])
def apply_promotion_code():
    if 'user_id' not in session or session.get('role') != 'customer':
        flash('You need to login to use the promo code.', 'danger')
        return redirect(url_for('checkout'))

    if 'applied_promotion' in session:
        flash('You can only apply one promotion per order.', 'danger')
        return redirect(url_for('checkout'))
    
    promo_code = request.form.get('promo_code')
    order_total = float(request.form.get('order_total', 0))

    if not promo_code:
        flash('Please enter a promo code', 'danger')
        return redirect(url_for('checkout'))

    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("SELECT customer_id FROM Customers WHERE user_id = %s", (session['user_id'],))
        customer = cursor.fetchone()
        if not customer:
            flash('Cannot find customer information.', 'danger')
            return redirect(url_for('checkout'))
        
        customer_id = customer['customer_id']

        cursor.execute("""
            SELECT p.*
            FROM Promotions p
            WHERE p.promotion_code = %s
            AND NOW() BETWEEN p.start_date AND p.end_date
            AND (p.max_uses = 0 OR p.current_uses < p.max_uses)
            AND NOT EXISTS (
                SELECT 1 
                FROM Orders o 
                WHERE o.promotion_id = p.promotion_id
                AND o.customer_id = %s
                AND o.status != 'cancelled'
            )
        """, (promo_code, customer_id))
        promotion = cursor.fetchone()

        if not promotion:
            flash('Promotion code is invalid, expired, or already used', 'danger')
            return redirect(url_for('checkout'))

        min_order_amount = float(promotion.get('min_order_amount', 0))
        if order_total < min_order_amount:
            flash(f'Minimum order amount of {min_order_amount:,.0f} VND required for this promotion', 'danger')
            return redirect(url_for('checkout'))

        discount_percentage = float(promotion['discount_percentage'])
        discount_amount = order_total * (discount_percentage / 100)

        max_discount = float(promotion.get('max_discount_amount', float('inf')))
        discount_amount = min(discount_amount, max_discount)
        
        final_price = order_total - discount_amount

        session['applied_promotion'] = {
            'code': promo_code,
            'discount_percentage': discount_percentage,
            'discount_amount': round(discount_amount, 2),
            'final_price': round(final_price, 2),
            'promotion_id': promotion['promotion_id'],
            'min_order_amount': min_order_amount,
            'max_discount_amount': max_discount
        }

        flash(f'Successfully applied {discount_percentage}% discount!', 'success')
        return redirect(url_for('checkout'))

    except Exception as e:
        flash(f'Error applying promotion: {str(e)}', 'danger')
        return redirect(url_for('checkout'))
    finally:
        cursor.close()
  

@app.route('/remove_promotion_code', methods=['POST'])
def remove_promotion_code():
    if 'applied_promotion' in session:
        session.pop('applied_promotion', None)
        flash('Deleted the promotion code.', 'info')
    return redirect(url_for('checkout'))
                            
#------------------------------------------ADMIN: STATISTICS------------------------------------------------
@app.route('/admin_statistics', methods=['GET'])
def admin_statistics():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))
    
    cursor = mysql.connection.cursor(DictCursor)
    try:
        cursor.execute("""
            SELECT 
                YEAR(o.created_at) as year,
                MONTH(o.created_at) as month,
                COUNT(o.order_id) as total_orders,
                COALESCE(SUM(o.total_price), 0) as revenue,
                COUNT(CASE WHEN o.status = 'completed' THEN 1 END) as completed_orders,
                COUNT(CASE WHEN o.status IN ('processing', 'accepted', 'delivering') THEN 1 END) as processing_orders
            FROM Orders o
            GROUP BY YEAR(o.created_at), MONTH(o.created_at)
            ORDER BY year, month
        """)
        monthly_stats = cursor.fetchall()

        cursor.execute("""
            SELECT 
                COUNT(order_id) as total_orders,
                COALESCE(SUM(total_price), 0) as total_revenue,
                COUNT(CASE WHEN status = 'completed' THEN 1 END) as total_completed,
                COUNT(CASE WHEN status IN ('processing', 'accepted', 'delivering') THEN 1 END) as total_processing
            FROM Orders
        """)
        total_stats = cursor.fetchone()

        cursor.execute("""
            SELECT 
                p.product_name,
                COALESCE(SUM(oi.quantity), 0) as quantity_sold,
                COALESCE(SUM(oi.subtotal), 0) as total_revenue
            FROM Products p
            LEFT JOIN Order_Items oi ON p.product_id = oi.product_id
            LEFT JOIN Orders o ON oi.order_id = o.order_id
            GROUP BY p.product_id, p.product_name
            ORDER BY quantity_sold DESC
            LIMIT 10
        """)
        top_products = cursor.fetchall()
        
        return render_template('admin/admin_statistics.html', 
                            top_products=top_products,
                            monthly_stats=monthly_stats,
                            total_stats=total_stats)

    except Exception as e:
        flash(f'Error occurred: {str(e)}', 'danger')
        return redirect(url_for('admin_dashboard'))

    finally:
        if cursor:
            cursor.close()
            
@app.route('/generate_report', methods=['GET', 'POST'])
def generate_report():
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You need to login as administrator.', 'danger')
        return redirect(url_for('login'))

    try:
        doc = Document()

        title = doc.add_heading('Moonwalker Cake Shop - Statistics Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading('Monthly Statistics', level=2)

        monthly_table = doc.add_table(rows=1, cols=5)
        monthly_table.style = 'Table Grid'

        hdr_cells = monthly_table.rows[0].cells
        hdr_cells[0].text = 'Month'
        hdr_cells[1].text = 'Total Orders'
        hdr_cells[2].text = 'Revenue'
        hdr_cells[3].text = 'Completed Orders'
        hdr_cells[4].text = 'Processing Orders'

        cursor = mysql.connection.cursor(DictCursor)
        cursor.execute("""
            SELECT 
                YEAR(created_at) as year,
                MONTH(created_at) as month,
                COUNT(*) as total_orders,
                SUM(total_price) as revenue,
                SUM(CASE WHEN status = 'completed' THEN 1 ELSE 0 END) as completed_orders,
                SUM(CASE WHEN status = 'processing' THEN 1 ELSE 0 END) as processing_orders
            FROM Orders
            GROUP BY YEAR(created_at), MONTH(created_at)
            ORDER BY year, month
        """)
        monthly_stats = cursor.fetchall()

        for stat in monthly_stats:
            row_cells = monthly_table.add_row().cells
            row_cells[0].text = f"{stat['month']}/2024"
            row_cells[1].text = str(stat['total_orders'])
            row_cells[2].text = "{:,.0f} VNĐ".format(stat['revenue'])
            row_cells[3].text = str(stat['completed_orders'])
            row_cells[4].text = str(stat['processing_orders'])

        doc.add_heading('Top Selling Products', level=2)

        products_table = doc.add_table(rows=1, cols=3)
        products_table.style = 'Table Grid'

        hdr_cells = products_table.rows[0].cells
        hdr_cells[0].text = 'Product Name'
        hdr_cells[1].text = 'Quantity Sold'
        hdr_cells[2].text = 'Revenue'

        cursor.execute("""
            SELECT 
                p.product_name,
                SUM(oi.quantity) as quantity_sold,
                SUM(oi.subtotal) as total_revenue
            FROM Order_Items oi
            JOIN Products p ON oi.product_id = p.product_id
            JOIN Orders o ON oi.order_id = o.order_id
            GROUP BY p.product_id
            ORDER BY quantity_sold DESC
            LIMIT 10
        """)
        top_products = cursor.fetchall()
        
        for product in top_products:
            row_cells = products_table.add_row().cells
            row_cells[0].text = product['product_name']
            row_cells[1].text = str(product['quantity_sold'])
            row_cells[2].text = "{:,.0f} VNĐ".format(product['total_revenue'])
        
        doc.add_paragraph(f"Report generated on {datetime.now().strftime('%Y-%m-%d')}")
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        from flask import send_file
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"Moonwalker_Statistics_{datetime.now().strftime('%Y-%m-%d')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        flash(f'Error generating report: {str(e)}', 'danger')
        return redirect(url_for('admin_statistics'))
    finally:
        if cursor:
            cursor.close()
#-------------------------------------------MAIN GATEWAY------------------------------------------------
if not os.path.exists('static/uploads/payments'):
    os.makedirs('static/uploads/payments')

if __name__ == '__main__':
    app.run()
